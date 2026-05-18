"""
历史记录和导出路由
/api/history, /api/export/markdown, /api/video/generate
"""
import os
import io
import re
import logging
import zipfile

import requests as http_requests
from flask import Blueprint, Response, jsonify, request
from urllib.parse import urlparse

from services.database_service import get_db_service
from services.video_service import get_video_service
from services.oss_service import get_oss_service

logger = logging.getLogger(__name__)

history_bp = Blueprint('history', __name__)


@history_bp.route('/api/history', methods=['GET'])
def list_history():
    """获取历史记录列表（支持分页和类型筛选）"""
    try:
        page = request.args.get('page', 1, type=int)
        page_size = request.args.get('page_size', 12, type=int)
        content_type = request.args.get('type', 'all')
        offset = (page - 1) * page_size

        db_service = get_db_service()

        total = db_service.count_history_by_type(content_type if content_type != 'all' else None)
        records = db_service.list_history_by_type(
            content_type=content_type if content_type != 'all' else None,
            limit=page_size,
            offset=offset
        )
        total_pages = (total + page_size - 1) // page_size

        return jsonify({
            'success': True,
            'records': records,
            'total': total,
            'page': page,
            'page_size': page_size,
            'total_pages': total_pages,
            'content_type': content_type
        })
    except Exception as e:
        logger.error(f"获取历史记录失败: {e}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)}), 500


@history_bp.route('/api/history/<history_id>', methods=['GET'])
def get_history(history_id):
    """获取单条历史记录详情"""
    try:
        db_service = get_db_service()
        record = db_service.get_history(history_id)
        if record:
            return jsonify({'success': True, 'record': record})
        else:
            return jsonify({'success': False, 'error': '记录不存在'}), 404
    except Exception as e:
        logger.error(f"获取历史记录失败: {e}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)}), 500


@history_bp.route('/api/history/<history_id>', methods=['DELETE'])
def delete_history(history_id):
    """删除历史记录"""
    try:
        db_service = get_db_service()
        deleted = db_service.delete_history(history_id)
        if deleted:
            return jsonify({'success': True, 'message': '删除成功'})
        else:
            return jsonify({'success': False, 'error': '记录不存在'}), 404
    except Exception as e:
        logger.error(f"删除历史记录失败: {e}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)}), 500


@history_bp.route('/api/video/generate', methods=['POST'])
def generate_video():
    """生成封面动画视频"""
    try:
        data = request.get_json()

        if not data:
            return jsonify({'success': False, 'error': '请提供 JSON 数据'}), 400

        history_id = data.get('history_id')
        image_url = data.get('image_url')
        image_path = data.get('image_path')
        prompt = data.get('prompt')

        video_service = get_video_service()
        if not video_service or not video_service.is_available():
            return jsonify({'success': False, 'error': '视频生成服务不可用'}), 503

        if not image_url:
            if not image_path:
                if history_id:
                    db_service = get_db_service()
                    record = db_service.get_history(history_id)
                    if record and record.get('cover_image'):
                        image_path = record.get('cover_image')

            if not image_path:
                return jsonify({'success': False, 'error': '缺少 image_url 或 image_path 参数'}), 400

            oss_service = get_oss_service()
            if not oss_service or not oss_service.is_available:
                return jsonify({'success': False, 'error': 'OSS 服务不可用，无法上传图片'}), 503

            import uuid
            unique_id = uuid.uuid4().hex[:8]
            filename = os.path.basename(image_path)
            remote_path = f"vibe-blog/covers/{unique_id}_{filename}"

            oss_result = oss_service.upload_file(
                local_path=image_path,
                remote_path=remote_path
            )

            if not oss_result.get('success'):
                return jsonify({'success': False, 'error': f"图片上传失败: {oss_result.get('error')}"}), 500

            image_url = oss_result['url']
            logger.info(f"封面图已上传到 OSS: {image_url}")

        logger.info(f"开始生成封面动画: history_id={history_id}, image_url={image_url[:80]}...")

        result = video_service.generate_from_image(
            image_url=image_url,
            prompt=prompt
        )

        if not result:
            return jsonify({'success': False, 'error': '视频生成失败'}), 500

        video_filename = os.path.basename(result.local_path) if result.local_path else None
        video_access_url = f"/outputs/videos/{video_filename}" if video_filename else result.url

        if history_id:
            db_service = get_db_service()
            db_service.update_history_video(history_id, video_access_url)

        logger.info(f"封面动画生成成功: {video_access_url}")

        return jsonify({
            'success': True,
            'video_url': video_access_url,
            'task_id': result.task_id
        })

    except Exception as e:
        logger.error(f"视频生成失败: {e}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)}), 500


def _extract_image_urls(markdown_content):
    """从 Markdown 中提取所有图片 URL"""
    pattern = r'!\[([^\]]*)\]\(([^\)]+)\)'
    matches = re.findall(pattern, markdown_content)
    return matches


def _download_image(url, timeout=10):
    """下载图片，返回二进制内容"""
    try:
        original_url = url

        if url.startswith('./images/'):
            url = '/outputs/images/' + url[9:]
        elif url.startswith('/outputs/images/'):
            pass

        if url.startswith('/'):
            base_url = request.host_url.rstrip('/')
            url = base_url + url

        logger.info(f"下载图片: {original_url} -> {url}")
        response = http_requests.get(url, timeout=timeout, allow_redirects=True)
        response.raise_for_status()
        return response.content
    except Exception as e:
        logger.warning(f"下载图片失败 {original_url}: {e}")
        return None


def _get_image_filename(url):
    """从 URL 中提取文件名"""
    parsed = urlparse(url)
    filename = os.path.basename(parsed.path)
    if not filename or '.' not in filename:
        filename = 'image.png'
    return filename


@history_bp.route('/api/export/markdown', methods=['POST'])
def export_markdown_with_images():
    """导出 Markdown 文件，包含所有本地图片"""
    try:
        data = request.get_json()
        if not data or 'markdown' not in data:
            return jsonify({'success': False, 'error': '缺少 markdown 参数'}), 400

        markdown_content = data.get('markdown', '')
        title = data.get('title', 'blog')

        safe_title = re.sub(r'[^\w\u4e00-\u9fa5_-]', '_', title)[:50]

        image_matches = _extract_image_urls(markdown_content)

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            zip_file.comment = b''
            modified_markdown = markdown_content
            image_mapping = {}

            for alt_text, img_url in image_matches:
                img_content = _download_image(img_url)
                if img_content:
                    original_filename = _get_image_filename(img_url)
                    base_name, ext = os.path.splitext(original_filename)
                    counter = 1
                    new_filename = original_filename
                    while new_filename in image_mapping.values():
                        new_filename = f"{base_name}_{counter}{ext}"
                        counter += 1

                    zip_file.writestr(f'images/{new_filename}', img_content)
                    image_mapping[img_url] = new_filename

                    old_ref = f'![{alt_text}]({img_url})'
                    new_ref = f'![{alt_text}](./images/{new_filename})'
                    modified_markdown = modified_markdown.replace(old_ref, new_ref)

            zip_file.writestr(f'{safe_title}.md', modified_markdown.encode('utf-8'))

        zip_buffer.seek(0)
        timestamp = __import__('datetime').datetime.now().strftime('%Y%m%d')
        filename = f'export_{timestamp}.zip'

        return Response(
            zip_buffer.getvalue(),
            mimetype='application/zip',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"'
            }
        )

    except Exception as e:
        logger.error(f"导出 Markdown 失败: {e}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)}), 500


@history_bp.route('/api/export/word', methods=['POST'])
def export_word():
    """导出 Word (.docx) 文件"""
    try:
        data = request.get_json()
        if not data or 'markdown' not in data:
            return jsonify({'success': False, 'error': '缺少 markdown 参数'}), 400

        markdown_content = data.get('markdown', '')
        title = data.get('title', 'blog')
        safe_title = re.sub(r'[^\w\u4e00-\u9fa5_-]', '_', title)[:50]

        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return jsonify({'success': False, 'error': '服务端未安装 python-docx，请运行 pip install python-docx'}), 500

        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(11)

        lines = markdown_content.split('\n')
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # 标题
            if stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('#### '):
                doc.add_heading(stripped[5:], level=4)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            elif re.match(r'^\d+\.\s', stripped):
                text = re.sub(r'^\d+\.\s', '', stripped)
                doc.add_paragraph(text, style='List Number')
            elif stripped.startswith('> '):
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                p.paragraph_format.left_indent = Inches(0.5)
                # 去除 markdown 加粗/斜体标记
                clean = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped[2:])
                clean = re.sub(r'\*(.*?)\*', r'\1', clean)
                p.add_run(clean)
            elif stripped.startswith('```'):
                # 跳过代码块标记行
                continue
            elif stripped.startswith('!['):
                # 跳过图片标记
                continue
            else:
                # 普通段落：去除 markdown 内联标记
                clean = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
                clean = re.sub(r'\*(.*?)\*', r'\1', clean)
                clean = re.sub(r'`([^`]+)`', r'\1', clean)
                clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', clean)
                doc.add_paragraph(clean)

        # 保存到内存
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return Response(
            buffer.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f'attachment; filename="{safe_title}.docx"'
            }
        )

    except Exception as e:
        logger.error(f"导出 Word 失败: {e}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)}), 500
